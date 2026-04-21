"""Minimal stdlib-only tests for the Input Expansion module.

Run directly:

    python3 tests/test_parse_inputs.py

Covered:
- unsupported file type -> failed_sources with reason=unsupported_file_type
- empty text file        -> empty_extracted_sources
- txt / md               -> documents
- docx                   -> documents (requires python-docx)
- image w/o OCR backend  -> failed_sources with reason=ocr_backend_unavailable
- notifier event stream  -> detected / start / success|failed / summary

NOTE: This script is self-contained and does NOT require pytest.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
from pathlib import Path

# Make the project importable when running `python3 tests/test_parse_inputs.py`.
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from tools.parse_inputs import (  # noqa: E402
    FILE_NOT_FOUND,
    IMAGE_EXTENSIONS,
    OCR_BACKEND_UNAVAILABLE,
    PARSE_ERROR,
    SUPPORTED_EXTENSIONS,
    UNSUPPORTED_FILE_TYPE,
    parse_inputs,
)
import tools.parse_inputs as pi  # noqa: E402


_passed = 0
_failed = 0


def _check(name, cond, detail=""):
    global _passed, _failed
    if cond:
        _passed += 1
        print(f"  PASS: {name}")
    else:
        _failed += 1
        print(f"  FAIL: {name} -- {detail}")


def test_supported_extensions_declared():
    print("[test] supported extension surface")
    for ext in (".txt", ".md", ".pdf", ".docx", ".png", ".jpg", ".jpeg"):
        _check(f"{ext} declared", ext in SUPPORTED_EXTENSIONS)


def test_unsupported_file_type():
    print("[test] unsupported file type")
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "foo.log"
        p.write_text("noop", encoding="utf-8")
        result = parse_inputs([str(p)])
        failed = result["logs"]["failed_sources"]
        _check("one failed entry", len(failed) == 1, str(failed))
        _check(
            "reason tag",
            failed and failed[0]["reason"] == UNSUPPORTED_FILE_TYPE,
            str(failed),
        )
        _check(
            "source_type carried",
            failed and failed[0]["source_type"] == "log",
            str(failed),
        )


def test_txt_and_empty():
    print("[test] txt ok + empty txt")
    with tempfile.TemporaryDirectory() as d:
        good = Path(d) / "good.txt"
        good.write_text("hello world", encoding="utf-8")
        empty = Path(d) / "empty.txt"
        empty.write_text("   \n\n  ", encoding="utf-8")
        result = parse_inputs([str(good), str(empty)])
        _check("two documents", len(result["documents"]) == 2)
        empties = result["logs"]["empty_extracted_sources"]
        _check("empty recorded", "empty.txt" in empties, str(empties))
        summary = result["ingestion_summary"]
        _check("summary detected", summary["detected"] == 2)
        _check("summary succeeded=1", summary["succeeded"] == 1, str(summary))
        _check("summary empty=1", summary["empty_extracted"] == 1, str(summary))


def test_missing_file():
    print("[test] file not found")
    result = parse_inputs(["/no/such/file.md"])
    failed = result["logs"]["failed_sources"]
    _check("one failed", len(failed) == 1)
    _check(
        "reason=file_not_found",
        failed and failed[0]["reason"] == FILE_NOT_FOUND,
        str(failed),
    )


def test_docx_roundtrip():
    print("[test] docx happy path")
    try:
        from docx import Document
    except Exception as exc:
        _check("python-docx available", False, f"import failed: {exc}")
        return
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "note.docx"
        doc = Document()
        doc.add_paragraph("基础概念：监督学习是从标注数据学习映射。")
        doc.add_paragraph("方法步骤：清洗、切分、训练。")
        doc.save(str(p))
        result = parse_inputs([str(p)])
        docs = result["documents"]
        _check("docx parsed", len(docs) == 1, str(result))
        if docs:
            _check(
                "source_type=docx",
                docs[0]["source_type"] == "docx",
                str(docs[0]),
            )
            _check(
                "text extracted",
                "监督学习" in docs[0]["extracted_text"],
                docs[0]["extracted_text"][:80],
            )


def test_image_degrades_without_backend():
    print("[test] image without OCR backend degrades gracefully")
    # Create a tiny 1x1 file. We don't need a real image because probe fails
    # before the reader actually opens the file -- in this sandbox the
    # tesseract binary is missing, so the OCR probe must short-circuit.
    import shutil

    if shutil.which("tesseract"):
        print("  SKIP: tesseract binary is present; degrade path not applicable")
        return

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "fake.png"
        p.write_bytes(b"\x89PNG\r\n\x1a\n")  # just the PNG signature
        result = parse_inputs([str(p)])
        failed = result["logs"]["failed_sources"]
        _check("one failed", len(failed) == 1, str(failed))
        if failed:
            _check(
                "reason=ocr_backend_unavailable",
                failed[0]["reason"] == OCR_BACKEND_UNAVAILABLE,
                str(failed[0]),
            )


def test_notifier_event_stream():
    print("[test] notifier fires expected events")
    events = []

    def recorder(event, payload):
        events.append((event, payload))

    with tempfile.TemporaryDirectory() as d:
        good = Path(d) / "good.txt"
        good.write_text("hi", encoding="utf-8")
        bad = Path(d) / "weird.log"
        bad.write_text("x", encoding="utf-8")
        parse_inputs([str(good), str(bad)], notifier=recorder)

    names = [e[0] for e in events]
    _check("detected first", names[0] == "detected", str(names))
    _check("summary last", names[-1] == "summary", str(names))
    _check("has start events", names.count("start") == 2, str(names))
    _check("has one success", names.count("success") == 1, str(names))
    _check("has one failed", names.count("failed") == 1, str(names))


def test_pdf_encrypted_fails_gracefully():
    print("[test] encrypted pdf fails gracefully")
    try:
        from pypdf import PdfWriter
    except Exception as exc:
        print(f"  SKIP: pypdf unavailable ({exc})")
        return

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "encrypted.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.encrypt("secret")
        with p.open("wb") as fp:
            writer.write(fp)

        result = parse_inputs([str(p)])
        failed = result["logs"]["failed_sources"]
        _check("one failed", len(failed) == 1, str(failed))
        if failed:
            _check(
                "reason=parse_error",
                failed[0]["reason"] == PARSE_ERROR,
                str(failed[0]),
            )


def test_docx_heading_path_injected():
    print("[test] docx heading path injected")
    try:
        from docx import Document
    except Exception as exc:
        _check("python-docx available", False, f"import failed: {exc}")
        return

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "heading.docx"
        doc = Document()
        doc.add_heading("概念", level=1)
        doc.add_paragraph("监督学习是一类方法。")
        doc.save(str(p))

        result = parse_inputs([str(p)])
        docs = result["documents"]
        _check("docx parsed", len(docs) == 1, str(result))
        if docs:
            txt = docs[0]["extracted_text"]
            _check("contains heading path marker", "heading_path:" in txt, txt[:160])


def test_image_api_fallback_when_local_ocr_unavailable():
    print("[test] image api fallback when local OCR unavailable")
    old_probe = pi._probe_ocr_backend
    old_urlopen = pi.request.urlopen
    old_url = os.environ.get("IMAGE_OCR_API_URL")
    old_style = os.environ.get("IMAGE_OCR_API_STYLE")

    class FakeResp:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return json.dumps({"text": "这是图片 OCR API 提取文本"}).encode("utf-8")

    try:
        # Force local OCR to be unavailable, then ensure API fallback works.
        pi._probe_ocr_backend = lambda: (False, "missing local ocr backend")
        pi.request.urlopen = lambda req, timeout=0: FakeResp()
        os.environ["IMAGE_OCR_API_URL"] = "http://fake.local/image-ocr"
        os.environ["IMAGE_OCR_API_STYLE"] = "custom"

        with tempfile.TemporaryDirectory() as d:
            p = Path(d) / "img.png"
            p.write_bytes(b"\x89PNG\r\n\x1a\nfake")
            result = parse_inputs([str(p)], api_assist_enabled=True)

        docs = result["documents"]
        _check("image parsed via api fallback", len(docs) == 1, str(result))
        if docs:
            _check(
                "api text extracted",
                "OCR API 提取文本" in docs[0].get("extracted_text", ""),
                docs[0].get("extracted_text", ""),
            )
            _check(
                "backend marked as image_api_ocr",
                docs[0].get("extraction_backend") == "image_api_ocr",
                str(docs[0]),
            )
        summary = result.get("ingestion_summary", {})
        _check("image api attempted count", summary.get("image_api_attempted", 0) >= 1, str(summary))
        _check("image api succeeded count", summary.get("image_api_succeeded", 0) == 1, str(summary))
        _check("image api enhanced count", summary.get("image_api_enhanced", 0) == 0, str(summary))
    finally:
        pi._probe_ocr_backend = old_probe
        pi.request.urlopen = old_urlopen
        if old_url is None:
            os.environ.pop("IMAGE_OCR_API_URL", None)
        else:
            os.environ["IMAGE_OCR_API_URL"] = old_url
        if old_style is None:
            os.environ.pop("IMAGE_OCR_API_STYLE", None)
        else:
            os.environ["IMAGE_OCR_API_STYLE"] = old_style


def test_image_api_auto_enhance_replaces_low_quality_local():
    print("[test] image api auto enhance replaces low quality local")
    old_probe = pi._probe_ocr_backend
    old_read_image = pi._read_image_file
    old_call_api = pi._call_image_ocr_api
    old_url = os.environ.get("IMAGE_OCR_API_URL")
    old_style = os.environ.get("IMAGE_OCR_API_STYLE")
    try:
        pi._probe_ocr_backend = lambda: (True, "pytesseract")
        pi._read_image_file = lambda *args, **kwargs: "abc"  # low score
        pi._call_image_ocr_api = lambda *args, **kwargs: "这是 API 增强后的更完整图片文本内容"
        os.environ["IMAGE_OCR_API_URL"] = "http://fake.local/image-ocr"
        os.environ["IMAGE_OCR_API_STYLE"] = "custom"

        with tempfile.TemporaryDirectory() as d:
            p = Path(d) / "img.png"
            p.write_bytes(b"\x89PNG\r\n\x1a\nfake")
            result = parse_inputs(
                [str(p)],
                api_assist_enabled=True,
                image_api_enhance_mode="auto",
                image_api_enhance_min_score=20,
                image_api_enhance_ratio=1.05,
                image_api_enhance_min_delta=2,
            )
        docs = result["documents"]
        _check("image parsed", len(docs) == 1, str(result))
        if docs:
            _check("api used", docs[0].get("image_api_used") is True, str(docs[0]))
            _check(
                "backend enhanced tag",
                docs[0].get("extraction_backend") == "image_api_ocr_enhanced",
                str(docs[0]),
            )
            _check("enhanced flag", docs[0].get("image_api_enhanced") is True, str(docs[0]))
        summary = result.get("ingestion_summary", {})
        _check("attempted>=1", summary.get("image_api_attempted", 0) >= 1, str(summary))
        _check("succeeded=1", summary.get("image_api_succeeded", 0) == 1, str(summary))
        _check("enhanced=1", summary.get("image_api_enhanced", 0) == 1, str(summary))
    finally:
        pi._probe_ocr_backend = old_probe
        pi._read_image_file = old_read_image
        pi._call_image_ocr_api = old_call_api
        if old_url is None:
            os.environ.pop("IMAGE_OCR_API_URL", None)
        else:
            os.environ["IMAGE_OCR_API_URL"] = old_url
        if old_style is None:
            os.environ.pop("IMAGE_OCR_API_STYLE", None)
        else:
            os.environ["IMAGE_OCR_API_STYLE"] = old_style


def test_image_api_auto_enhance_keeps_better_local_text():
    print("[test] image api auto enhance keeps better local text")
    old_probe = pi._probe_ocr_backend
    old_read_image = pi._read_image_file
    old_call_api = pi._call_image_ocr_api
    old_url = os.environ.get("IMAGE_OCR_API_URL")
    old_style = os.environ.get("IMAGE_OCR_API_STYLE")
    try:
        pi._probe_ocr_backend = lambda: (True, "pytesseract")
        pi._read_image_file = (
            lambda *args, **kwargs: "本地 OCR 已经提取到了足够详细且可读的文本内容用于学习整理"
        )
        pi._call_image_ocr_api = lambda *args, **kwargs: "短文本"
        os.environ["IMAGE_OCR_API_URL"] = "http://fake.local/image-ocr"
        os.environ["IMAGE_OCR_API_STYLE"] = "custom"

        with tempfile.TemporaryDirectory() as d:
            p = Path(d) / "img.png"
            p.write_bytes(b"\x89PNG\r\n\x1a\nfake")
            result = parse_inputs(
                [str(p)],
                api_assist_enabled=True,
                image_api_enhance_mode="auto",
                image_api_enhance_min_score=999,  # force API attempt for comparison
                image_api_enhance_ratio=1.2,
                image_api_enhance_min_delta=10,
            )
        docs = result["documents"]
        _check("image parsed", len(docs) == 1, str(result))
        if docs:
            _check("keep local backend", docs[0].get("extraction_backend") == "pytesseract", str(docs[0]))
            _check("api not used", docs[0].get("image_api_used") is False, str(docs[0]))
            _check(
                "local text preserved",
                "本地 OCR 已经提取" in docs[0].get("extracted_text", ""),
                docs[0].get("extracted_text", ""),
            )
        summary = result.get("ingestion_summary", {})
        _check("attempted>=1", summary.get("image_api_attempted", 0) >= 1, str(summary))
        _check("succeeded=0", summary.get("image_api_succeeded", 0) == 0, str(summary))
    finally:
        pi._probe_ocr_backend = old_probe
        pi._read_image_file = old_read_image
        pi._call_image_ocr_api = old_call_api
        if old_url is None:
            os.environ.pop("IMAGE_OCR_API_URL", None)
        else:
            os.environ["IMAGE_OCR_API_URL"] = old_url
        if old_style is None:
            os.environ.pop("IMAGE_OCR_API_STYLE", None)
        else:
            os.environ["IMAGE_OCR_API_STYLE"] = old_style


def main():
    print("=" * 60)
    print("Input Expansion + Ingestion Notice: minimal tests")
    print("=" * 60)
    test_supported_extensions_declared()
    test_unsupported_file_type()
    test_txt_and_empty()
    test_missing_file()
    test_docx_roundtrip()
    test_image_degrades_without_backend()
    test_notifier_event_stream()
    test_pdf_encrypted_fails_gracefully()
    test_docx_heading_path_injected()
    test_image_api_fallback_when_local_ocr_unavailable()
    test_image_api_auto_enhance_replaces_low_quality_local()
    test_image_api_auto_enhance_keeps_better_local_text()
    print("-" * 60)
    print(f"Result: {_passed} passed, {_failed} failed")
    sys.exit(0 if _failed == 0 else 1)


if __name__ == "__main__":
    main()
