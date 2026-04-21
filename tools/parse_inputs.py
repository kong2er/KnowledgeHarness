"""Input parsing utilities for KnowledgeHarness.

Supported input matrix (see docs/PROJECT_STATE.md §2):

- txt / md           : stdlib only
- pdf                : requires `pypdf` (declared in requirements.txt)
- docx               : requires `python-docx` (declared in requirements.txt)
- png / jpg / jpeg   : opt-in OCR via `pytesseract` + `Pillow` + tesseract
                       binary. When any of those is missing the parser
                       degrades gracefully with reason=ocr_backend_unavailable
                       rather than silently pretending success.
                       When API assist is explicitly enabled and API is
                       configured, images can additionally use API OCR as
                       fallback/enhancement.

Failure semantics (see docs/ACCEPTANCE.md §4 `parse_inputs`):

    failed_sources entry = {
        "source": str(path),
        "source_name": basename,
        "source_type": ext without leading dot,
        "reason": one of UNSUPPORTED_FILE_TYPE | FILE_NOT_FOUND
                       | PARSE_ERROR | OCR_BACKEND_UNAVAILABLE,
        "error": human-readable detail,
    }

A per-file ``notifier`` callback is available so the CLI layer can surface
real-time ingestion progress without coupling this module to stdout.
"""

from __future__ import annotations

import base64
import json
import os
import shutil
import re
import mimetypes
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Tuple
from urllib import parse, request

# --- Extension sets -----------------------------------------------------

TEXT_EXTENSIONS: frozenset = frozenset({".txt", ".md"})
PDF_EXTENSIONS: frozenset = frozenset({".pdf"})
DOCX_EXTENSIONS: frozenset = frozenset({".docx"})
IMAGE_EXTENSIONS: frozenset = frozenset({".png", ".jpg", ".jpeg"})

# Full declared support surface (what app.py should glob for in dir mode).
SUPPORTED_EXTENSIONS: frozenset = (
    TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | IMAGE_EXTENSIONS
)

# --- Failure reason constants ------------------------------------------

UNSUPPORTED_FILE_TYPE = "unsupported_file_type"
FILE_NOT_FOUND = "file_not_found"
PARSE_ERROR = "parse_error"
OCR_BACKEND_UNAVAILABLE = "ocr_backend_unavailable"
DEFAULT_API_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "config" / "api_payload_templates.json"
)

# --- Notifier callback type --------------------------------------------

Notifier = Callable[[str, Dict[str, Any]], None]


def _emit(notifier: Optional[Notifier], event: str, payload: Dict[str, Any]) -> None:
    if notifier is None:
        return
    try:
        notifier(event, payload)
    except Exception:
        # Notifier must never break the pipeline.
        pass


# --- Custom exceptions --------------------------------------------------


class UnsupportedFileType(ValueError):
    """Raised when the file extension is not in SUPPORTED_EXTENSIONS."""


class OCRBackendUnavailable(RuntimeError):
    """Raised when pytesseract / Pillow / tesseract binary is missing."""

    def __init__(self, message: str, image_api_attempts: int = 0):
        super().__init__(message)
        self.image_api_attempts = int(image_api_attempts)


# --- Readers ------------------------------------------------------------


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf_file(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency probe
        raise RuntimeError(
            "PDF parsing requires 'pypdf'. Run: pip install -r requirements.txt"
        ) from exc

    reader = PdfReader(str(path))
    pages: List[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages).strip()


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency probe
        raise RuntimeError(
            "DOCX parsing requires 'python-docx'. "
            "Run: pip install -r requirements.txt"
        ) from exc

    document = Document(str(path))
    parts: List[str] = []
    heading_stack: Dict[int, str] = {}

    def _heading_level(style_name: str) -> Optional[int]:
        raw = (style_name or "").strip()
        s = raw.lower()
        # English style variants: Heading 1 / Heading1 / heading 2 ...
        m = re.match(r"heading\s*(\d*)", s)
        if m:
            num = m.group(1)
            return int(num) if num else 1
        # Chinese localized style variants, e.g. 标题 1 / 标题1
        if "标题" in raw:
            m = re.search(r"标题\s*(\d*)", raw)
            if m:
                num = m.group(1)
                return int(num) if num else 1
        return None

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        level = _heading_level(getattr(getattr(para, "style", None), "name", ""))
        if level is not None:
            heading_stack[level] = text
            for k in list(heading_stack.keys()):
                if k > level:
                    del heading_stack[k]
            parts.append(text)
            continue

        if heading_stack:
            path_text = " > ".join(heading_stack[k] for k in sorted(heading_stack.keys()))
            parts.append(f"{text} [heading_path: {path_text}]")
        else:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            joined = " | ".join(c for c in cells if c)
            if joined:
                parts.append(joined)

    return "\n\n".join(parts).strip()


# OCR backend probe result is cached per-process; the check is cheap but
# doing it once avoids repeating the same warning for every image file.
_OCR_PROBE_CACHE: Optional[Tuple[bool, str]] = None


def _probe_ocr_backend() -> Tuple[bool, str]:
    """Return (available, detail_message)."""
    global _OCR_PROBE_CACHE
    if _OCR_PROBE_CACHE is not None:
        return _OCR_PROBE_CACHE

    try:
        import pytesseract  # type: ignore  # noqa: F401
        from PIL import Image  # type: ignore  # noqa: F401
    except Exception as exc:
        _OCR_PROBE_CACHE = (
            False,
            f"python packages missing (pytesseract/Pillow): {exc}",
        )
        return _OCR_PROBE_CACHE

    if not shutil.which("tesseract"):
        _OCR_PROBE_CACHE = (
            False,
            "tesseract system binary not found on PATH; "
            "install it via your OS package manager "
            "(e.g. `apt-get install tesseract-ocr`).",
        )
        return _OCR_PROBE_CACHE

    try:
        import pytesseract  # type: ignore

        pytesseract.get_tesseract_version()
    except Exception as exc:
        _OCR_PROBE_CACHE = (False, f"tesseract invocation failed: {exc}")
        return _OCR_PROBE_CACHE

    _OCR_PROBE_CACHE = (True, "pytesseract")
    return _OCR_PROBE_CACHE


def _read_image_file(
    path: Path,
    ocr_languages: str = "chi_sim+eng",
    ocr_fallback_language: str = "eng",
) -> str:
    available, detail = _probe_ocr_backend()
    if not available:
        raise OCRBackendUnavailable(detail)

    from PIL import Image  # type: ignore
    import pytesseract  # type: ignore

    with Image.open(str(path)) as img:
        # Prefer Chinese+English joint OCR; fall back to English-only when
        # the chi_sim language pack is unavailable.
        try:
            text = pytesseract.image_to_string(img, lang=ocr_languages)
        except pytesseract.TesseractError:
            text = pytesseract.image_to_string(img, lang=ocr_fallback_language)

    return (text or "").strip()


def _load_image_ocr_api_template() -> Dict[str, Any]:
    template_path = Path(
        os.getenv("IMAGE_OCR_API_TEMPLATE", str(DEFAULT_API_TEMPLATE_PATH))
    )
    try:
        raw = json.loads(template_path.read_text(encoding="utf-8"))
        section = raw.get("image_ocr", {}) if isinstance(raw, dict) else {}
        return section if isinstance(section, dict) else {}
    except Exception:
        return {}


def _extract_json_object_from_text(content: str) -> Dict[str, Any]:
    text = (content or "").strip()
    if not text:
        raise ValueError("api returned empty text content")
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9_-]*\n?", "", text).strip()
        text = re.sub(r"\n?```$", "", text).strip()
    try:
        data = json.loads(text)
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        snippet = text[start : end + 1]
        data = json.loads(snippet)
        if isinstance(data, dict):
            return data
    raise ValueError("api text content is not valid JSON object")


def _resolve_api_style(url: str, module_style_key: str) -> str:
    style = (
        os.getenv(module_style_key, "").strip().lower()
        or os.getenv("KNOWLEDGEHARNESS_API_STYLE", "").strip().lower()
        or "auto"
    )
    if style in {"custom", "openai_compatible"}:
        return style
    parsed = parse.urlparse(url)
    host = (parsed.netloc or "").lower()
    path = parsed.path or ""
    if "api.deepseek.com" in host or "api.openai.com" in host:
        return "openai_compatible"
    if path in {"", "/"}:
        return "openai_compatible"
    return "custom"


def _resolve_openai_endpoint(url: str) -> str:
    parsed = parse.urlparse(url)
    path = (parsed.path or "").rstrip("/")
    if path.endswith("/chat/completions"):
        final_path = path
    elif path.endswith("/v1"):
        final_path = f"{path}/chat/completions"
    elif path in {"", "/"}:
        final_path = "/v1/chat/completions"
    else:
        final_path = f"{path}/chat/completions"
    return parse.urlunparse(
        (parsed.scheme, parsed.netloc, final_path, parsed.params, parsed.query, parsed.fragment)
    )


def _is_image_api_configured() -> bool:
    return bool(
        os.getenv("IMAGE_OCR_API_URL", "").strip()
        or os.getenv("KNOWLEDGEHARNESS_API_URL", "").strip()
    )


def _score_ocr_text(text: str) -> int:
    """Score OCR quality with a simple "meaningful char count" heuristic."""
    s = (text or "").strip()
    if not s:
        return 0
    meaningful = sum(1 for ch in s if ch.isalnum() or "\u4e00" <= ch <= "\u9fff")
    return meaningful


def _call_image_ocr_api(path: Path, timeout_sec: float = 8.0) -> str:
    """Extract image text using optional external API.

    Env priority:
    - IMAGE_OCR_API_URL / KEY / STYLE / MODEL
    - fallback to KNOWLEDGEHARNESS_API_URL / KEY / STYLE / MODEL
    """
    url = (
        os.getenv("IMAGE_OCR_API_URL", "").strip()
        or os.getenv("KNOWLEDGEHARNESS_API_URL", "").strip()
    )
    if not url:
        raise RuntimeError("IMAGE_OCR_API_URL is not configured")

    raw_bytes = path.read_bytes()
    max_bytes = int(os.getenv("IMAGE_OCR_MAX_BYTES", str(8 * 1024 * 1024)))
    if len(raw_bytes) > max_bytes:
        raise ValueError(
            f"image too large for api ocr ({len(raw_bytes)} bytes > {max_bytes})"
        )
    mime = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    image_base64 = base64.b64encode(raw_bytes).decode("ascii")

    template = _load_image_ocr_api_template()
    system_prompt = str(template.get("system_prompt") or "").strip() or (
        "你是受约束的 OCR 提取器。请只提取图片中的文本，不要编造。"
    )
    output_contract = (
        template.get("output_contract")
        if isinstance(template.get("output_contract"), dict)
        else {"text": "string"}
    )

    headers = {"Content-Type": "application/json"}
    api_key = (
        os.getenv("IMAGE_OCR_API_KEY", "").strip()
        or os.getenv("KNOWLEDGEHARNESS_API_KEY", "").strip()
    )
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    style = _resolve_api_style(url, "IMAGE_OCR_API_STYLE")
    request_url = url
    request_payload: Dict[str, Any]
    if style == "openai_compatible":
        request_url = _resolve_openai_endpoint(url)
        model = (
            os.getenv("IMAGE_OCR_API_MODEL", "").strip()
            or os.getenv("KNOWLEDGEHARNESS_API_MODEL", "").strip()
            or "gpt-4o-mini"
        )
        request_payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": system_prompt + " 只允许返回 JSON 对象，字段为 text。",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "请提取这张图片中的文字，保留原始语言与换行。"},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime};base64,{image_base64}",
                            },
                        },
                    ],
                },
            ],
            "temperature": 0.0,
            "response_format": {"type": "json_object"},
        }
    else:
        request_payload = {
            "task": "image_ocr",
            "mime_type": mime,
            "image_base64": image_base64,
            "system_prompt": system_prompt,
            "output_contract": output_contract,
            "rules": {"extract_only": True, "do_not_invent_text": True},
        }

    req = request.Request(
        url=request_url,
        method="POST",
        data=json.dumps(request_payload).encode("utf-8"),
        headers=headers,
    )
    with request.urlopen(req, timeout=timeout_sec) as resp:
        body = resp.read().decode("utf-8", errors="replace")
    data = json.loads(body)
    if style == "openai_compatible":
        choices = data.get("choices", []) if isinstance(data, dict) else []
        if not choices:
            raise ValueError("openai-compatible api returned empty choices")
        content = str(((choices[0] or {}).get("message") or {}).get("content") or "")
        data = _extract_json_object_from_text(content)
    if not isinstance(data, dict):
        raise ValueError("image ocr api returned invalid payload")
    text = str(
        data.get("text") or data.get("extracted_text") or data.get("content") or ""
    ).strip()
    if not text:
        raise ValueError("image ocr api returned empty text")
    return text


# --- Per-file dispatch --------------------------------------------------


def parse_single_file(
    path: str | Path,
    ocr_languages: str = "chi_sim+eng",
    ocr_fallback_language: str = "eng",
    api_assist_enabled: bool = False,
    image_api_timeout_sec: float = 8.0,
    image_api_retries: int = 1,
    image_api_enhance_mode: str = "auto",
    image_api_enhance_min_score: int = 40,
    image_api_enhance_ratio: float = 1.1,
    image_api_enhance_min_delta: int = 6,
) -> Dict[str, Any]:
    """Parse one file into a normalized document object.

    Raises:
        FileNotFoundError: path does not point to an existing file.
        UnsupportedFileType: extension not in SUPPORTED_EXTENSIONS.
        OCRBackendUnavailable: image file, but OCR backend missing.
        Exception: any other parse-level failure from the underlying reader.
    """
    file_path = Path(path)
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"Input file not found: {file_path}")

    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(f"Unsupported file type: {file_path.name}")

    extraction_backend = ""
    image_api_used = False
    image_api_attempts = 0
    image_api_enhanced = False

    if ext in TEXT_EXTENSIONS:
        extracted = _read_text_file(file_path)
        extraction_backend = "text_reader"
    elif ext in PDF_EXTENSIONS:
        extracted = _read_pdf_file(file_path)
        extraction_backend = "pdf_reader"
    elif ext in DOCX_EXTENSIONS:
        extracted = _read_docx_file(file_path)
        extraction_backend = "docx_reader"
    elif ext in IMAGE_EXTENSIONS:
        local_error: Exception | None = None
        local_text = ""
        try:
            local_text = _read_image_file(
                file_path,
                ocr_languages=ocr_languages,
                ocr_fallback_language=ocr_fallback_language,
            )
            extraction_backend = "pytesseract"
        except Exception as exc:
            local_error = exc
        extracted = (local_text or "").strip()
        local_score = _score_ocr_text(extracted)

        mode = (image_api_enhance_mode or "auto").strip().lower()
        if mode not in {"fallback_only", "auto", "prefer_api"}:
            mode = "auto"

        should_try_api = bool(
            api_assist_enabled
            and (
                local_error is not None
                or not (extracted or "").strip()
                or mode == "prefer_api"
                or (mode == "auto" and local_score < int(image_api_enhance_min_score))
            )
            and _is_image_api_configured()
        )
        if should_try_api:
            retries = max(0, int(image_api_retries))
            last_exc: Exception | None = None
            for _ in range(retries + 1):
                image_api_attempts += 1
                try:
                    api_text = _call_image_ocr_api(
                        file_path,
                        timeout_sec=image_api_timeout_sec,
                    )
                    api_text = (api_text or "").strip()
                    if api_text:
                        api_score = _score_ocr_text(api_text)
                        use_api = False
                        if local_error is not None or not extracted:
                            use_api = True
                        elif mode == "prefer_api":
                            use_api = True
                        elif mode == "auto":
                            threshold = max(
                                local_score + int(image_api_enhance_min_delta),
                                int(local_score * float(image_api_enhance_ratio)),
                            )
                            if api_score >= threshold:
                                use_api = True
                        # fallback_only only uses API when local path failed/empty.
                        if use_api:
                            image_api_used = True
                            image_api_enhanced = bool(extracted)
                            extracted = api_text
                            extraction_backend = (
                                "image_api_ocr_enhanced"
                                if image_api_enhanced
                                else "image_api_ocr"
                            )
                    break
                except Exception as exc:
                    last_exc = exc
                    continue
            if not image_api_used and local_error is not None:
                if isinstance(local_error, OCRBackendUnavailable):
                    raise OCRBackendUnavailable(
                        f"{local_error}; image api fallback failed: {last_exc}",
                        image_api_attempts=image_api_attempts,
                    )
                raise local_error

        if local_error is not None and not image_api_used:
            if isinstance(local_error, OCRBackendUnavailable):
                raise OCRBackendUnavailable(
                    str(local_error),
                    image_api_attempts=image_api_attempts,
                )
            raise local_error
    else:  # pragma: no cover - SUPPORTED_EXTENSIONS covers all branches
        raise UnsupportedFileType(f"Unsupported file type: {file_path.name}")

    return {
        "source_name": file_path.name,
        "source_path": str(file_path.resolve()),
        "source_type": ext.lstrip("."),
        "chunk_id": None,
        "raw_text": extracted,
        "extracted_text": extracted,
        "extraction_backend": extraction_backend,
        "image_api_used": image_api_used,
        "image_api_attempts": image_api_attempts,
        "image_api_enhanced": image_api_enhanced,
    }


# --- Batch driver -------------------------------------------------------


def _effective_supported_extensions(api_assist_enabled: bool = False) -> List[str]:
    """Return extensions we can actually fulfil in the current environment.

    - txt/md : always
    - pdf    : always (pypdf is in requirements.txt)
    - docx   : only when python-docx is importable
    - images : when local OCR backend probe passes, or when API assist is
               enabled and image OCR API is configured
    """
    eff = sorted(TEXT_EXTENSIONS | PDF_EXTENSIONS)
    try:
        import docx  # type: ignore  # noqa: F401

        eff.extend(sorted(DOCX_EXTENSIONS))
    except Exception:
        pass
    if _probe_ocr_backend()[0] or (api_assist_enabled and _is_image_api_configured()):
        eff.extend(sorted(IMAGE_EXTENSIONS))
    return sorted(set(eff))


def _build_failed_entry(
    path: str,
    ext: str,
    reason: str,
    error: str,
) -> Dict[str, Any]:
    name = Path(path).name
    return {
        "source": str(path),
        "source_name": name,
        "source_type": ext.lstrip(".") if ext else "",
        "reason": reason,
        "error": error,
    }


def parse_inputs(
    paths: Iterable[str | Path],
    notifier: Optional[Notifier] = None,
    ocr_languages: str = "chi_sim+eng",
    ocr_fallback_language: str = "eng",
    api_assist_enabled: bool = False,
    image_api_timeout_sec: float = 8.0,
    image_api_retries: int = 1,
    image_api_enhance_mode: str = "auto",
    image_api_enhance_min_score: int = 40,
    image_api_enhance_ratio: float = 1.1,
    image_api_enhance_min_delta: int = 6,
) -> Dict[str, Any]:
    """Parse multiple input files with optional progress notification.

    Returns:
        {
          "documents": [ ... ],
          "logs": {
            "failed_sources": [ ... ],
            "empty_extracted_sources": [ ... ],
          },
          "ingestion_summary": {
            "detected": N,
            "supported": N,
            "unsupported": N,
            "succeeded": N,
            "empty_extracted": N,
            "failed": N,
            "breakdown_by_type": {"md": 1, "docx": 1, ...},
            "supported_extensions_effective": [".txt", ".md", ...],
            "image_extensions_opt_in": [".png", ".jpg", ".jpeg"],
            "ocr_backend": "pytesseract" | "unavailable",
          }
        }
    """
    path_list: List[str] = [str(p) for p in paths]
    documents: List[Dict[str, Any]] = []
    failed_sources: List[Dict[str, Any]] = []
    empty_extracted_sources: List[str] = []
    breakdown: Dict[str, int] = {}

    ocr_available, _ = _probe_ocr_backend()
    effective_exts = _effective_supported_extensions(api_assist_enabled=api_assist_enabled)

    _emit(
        notifier,
        "detected",
        {
            "count": len(path_list),
            "supported_extensions_effective": effective_exts,
            "ocr_backend": "pytesseract" if ocr_available else "unavailable",
        },
    )

    supported_count = 0
    unsupported_count = 0
    image_api_used_count = 0
    image_api_attempted_count = 0
    image_api_enhanced_count = 0

    for path in path_list:
        ext = Path(path).suffix.lower()
        name = Path(path).name
        is_supported = ext in SUPPORTED_EXTENSIONS
        if is_supported:
            supported_count += 1
        else:
            unsupported_count += 1

        breakdown[ext.lstrip(".") or "(noext)"] = (
            breakdown.get(ext.lstrip(".") or "(noext)", 0) + 1
        )

        _emit(
            notifier,
            "start",
            {
                "source_name": name,
                "source_path": path,
                "source_type": ext.lstrip("."),
                "supported": is_supported,
            },
        )

        try:
            doc = parse_single_file(
                path,
                ocr_languages=ocr_languages,
                ocr_fallback_language=ocr_fallback_language,
                api_assist_enabled=api_assist_enabled,
                image_api_timeout_sec=image_api_timeout_sec,
                image_api_retries=image_api_retries,
                image_api_enhance_mode=image_api_enhance_mode,
                image_api_enhance_min_score=image_api_enhance_min_score,
                image_api_enhance_ratio=image_api_enhance_ratio,
                image_api_enhance_min_delta=image_api_enhance_min_delta,
            )
        except UnsupportedFileType as exc:
            failed_sources.append(
                _build_failed_entry(path, ext, UNSUPPORTED_FILE_TYPE, str(exc))
            )
            _emit(
                notifier,
                "failed",
                {
                    "source_name": name,
                    "reason": UNSUPPORTED_FILE_TYPE,
                    "error": str(exc),
                },
            )
            continue
        except FileNotFoundError as exc:
            failed_sources.append(
                _build_failed_entry(path, ext, FILE_NOT_FOUND, str(exc))
            )
            _emit(
                notifier,
                "failed",
                {
                    "source_name": name,
                    "reason": FILE_NOT_FOUND,
                    "error": str(exc),
                },
            )
            continue
        except OCRBackendUnavailable as exc:
            image_api_attempted_count += int(getattr(exc, "image_api_attempts", 0) or 0)
            failed_sources.append(
                _build_failed_entry(
                    path, ext, OCR_BACKEND_UNAVAILABLE, str(exc)
                )
            )
            _emit(
                notifier,
                "failed",
                {
                    "source_name": name,
                    "reason": OCR_BACKEND_UNAVAILABLE,
                    "error": str(exc),
                },
            )
            continue
        except Exception as exc:  # any other parse-level failure
            failed_sources.append(
                _build_failed_entry(path, ext, PARSE_ERROR, str(exc))
            )
            _emit(
                notifier,
                "failed",
                {
                    "source_name": name,
                    "reason": PARSE_ERROR,
                    "error": str(exc),
                },
            )
            continue

        extracted = (doc.get("extracted_text") or "").strip()
        is_empty = not extracted
        if is_empty:
            empty_extracted_sources.append(doc.get("source_name") or path)

        documents.append(doc)
        image_api_attempted_count += int(doc.get("image_api_attempts") or 0)
        if bool(doc.get("image_api_used")):
            image_api_used_count += 1
        if bool(doc.get("image_api_enhanced")):
            image_api_enhanced_count += 1
        _emit(
            notifier,
            "success",
            {
                "source_name": doc.get("source_name"),
                "chars": len(extracted),
                "empty": is_empty,
            },
        )

    ingestion_summary = {
        "detected": len(path_list),
        "supported": supported_count,
        "unsupported": unsupported_count,
        "succeeded": len(documents) - len(empty_extracted_sources),
        "empty_extracted": len(empty_extracted_sources),
        "failed": len(failed_sources),
        "breakdown_by_type": breakdown,
        "supported_extensions_effective": effective_exts,
        "image_extensions_opt_in": sorted(IMAGE_EXTENSIONS),
        "ocr_backend": "pytesseract" if ocr_available else "unavailable",
        "image_api_assist_enabled": bool(api_assist_enabled),
        "image_api_enhance_mode": str(image_api_enhance_mode),
        "image_api_attempted": image_api_attempted_count,
        "image_api_succeeded": image_api_used_count,
        "image_api_enhanced": image_api_enhanced_count,
    }

    _emit(notifier, "summary", dict(ingestion_summary))

    return {
        "documents": documents,
        "logs": {
            "failed_sources": failed_sources,
            "empty_extracted_sources": empty_extracted_sources,
        },
        "ingestion_summary": ingestion_summary,
    }


if __name__ == "__main__":
    import json
    import sys

    parsed = parse_inputs(sys.argv[1:])
    print(json.dumps(parsed, ensure_ascii=False, indent=2))
