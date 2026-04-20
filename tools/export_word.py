"""Word (.docx) export from markdown output.

Scope: handles only the subset of Markdown that `tools/export_notes.
_render_final_notes_markdown` actually emits:

- ATX headings (`# / ## / ###`)
- bullet lines (`- `)
- blockquote lines (`> `)
- horizontal rules (`---` / `***` / `___`)
- inline italics via `*...*` → Word run with italic=True
- everything else → Normal paragraph

Not a general markdown-to-docx converter; do not use it as one.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

_HORIZONTAL_RULE_RE = re.compile(r"^(?:-{3,}|\*{3,}|_{3,})$")
# Non-greedy, no nested stars. Matches `*text*` but skips `**bold**` because
# the surrounding `*` greedily start a new group. For our current renderer
# that only emits single-star italics, this is enough.
_ITALIC_SPAN_RE = re.compile(r"\*([^*\n]+?)\*")


def _add_inline_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph`, turning `*...*` spans into italic runs."""
    if not text:
        return
    cursor = 0
    for match in _ITALIC_SPAN_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        italic_run = paragraph.add_run(match.group(1))
        italic_run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_styled_paragraph(doc, text: str, style: str | None = None):
    """Append a paragraph with the given style, gracefully falling back
    when the style is missing from the current template.
    """
    try:
        return doc.add_paragraph("", style=style) if style else doc.add_paragraph("")
    except KeyError:
        # Style not present in the default template -- use Normal.
        return doc.add_paragraph("")


def _markdown_line_to_paragraph(doc, line: str) -> None:
    s = line.rstrip()
    if not s:
        doc.add_paragraph("")
        return

    # Horizontal rule: python-docx has no native HR; a blank paragraph
    # reads cleaner than a literal "---".
    if _HORIZONTAL_RULE_RE.match(s):
        doc.add_paragraph("")
        return

    # Headings
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3)
        return
    if s.startswith("## "):
        doc.add_heading(s[3:].strip(), level=2)
        return
    if s.startswith("# "):
        doc.add_heading(s[2:].strip(), level=1)
        return

    # Blockquote → Quote style when available (falls back to Normal)
    if s.startswith("> "):
        body = s[2:].strip()
        paragraph = _add_styled_paragraph(doc, body, style="Quote")
        _add_inline_runs(paragraph, body)
        return

    # Bullet list item
    if s.startswith("- "):
        body = s[2:].strip()
        paragraph = _add_styled_paragraph(doc, body, style="List Bullet")
        _add_inline_runs(paragraph, body)
        return

    paragraph = doc.add_paragraph("")
    _add_inline_runs(paragraph, s)


def export_word_from_markdown(
    markdown_path: str | Path,
    out_dir: str | Path = "outputs",
    filename: str = "result.docx",
) -> str:
    """Generate a basic docx from markdown text.

    Requires `python-docx` (already part of project requirements).
    """
    from docx import Document  # lazy import

    md_path = Path(markdown_path)
    if not md_path.exists() or not md_path.is_file():
        raise FileNotFoundError(f"markdown file not found: {md_path}")

    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    docx_path = out / filename

    text = md_path.read_text(encoding="utf-8", errors="replace")
    lines: List[str] = text.splitlines()

    doc = Document()
    for line in lines:
        _markdown_line_to_paragraph(doc, line)
    doc.save(str(docx_path))
    return str(docx_path.resolve())

